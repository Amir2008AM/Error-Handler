#!/usr/bin/env python3
"""
PDF → Word converter
Primary:  pdfminer.six (text extraction) + python-docx (rebuild)
Fallback: LibreOffice headless
Output:   JSON to stdout { success, engine, pageCount, error }
"""
import sys, os, json, subprocess, tempfile
from pathlib import Path

def extract_with_pdfminer(pdf_path: str, docx_path: str) -> dict:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer, LTAnon, LTFigure
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()

    # Set A4 page
    from docx.util import Inches, Cm
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin   = Cm(2.5)
    section.top_margin  = section.bottom_margin  = Cm(2.5)

    page_count = 0
    has_arabic = False

    for page_layout in extract_pages(pdf_path):
        page_count += 1
        if page_count > 1:
            doc.add_page_break()

        for element in page_layout:
            if isinstance(element, LTTextContainer):
                raw = element.get_text().strip()
                if not raw:
                    continue

                # Detect Arabic text
                arabic_range = re.compile(r'[\u0600-\u06FF\u0750-\u077F]')
                is_arabic = bool(arabic_range.search(raw))
                if is_arabic:
                    has_arabic = True

                para = doc.add_paragraph()

                # Alignment
                if is_arabic:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    # Heuristic: if element is near center → center align
                    page_w = page_layout.width
                    el_cx  = element.x0 + (element.x1 - element.x0) / 2
                    if abs(el_cx - page_w / 2) < page_w * 0.1:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = para.add_run(raw)

                # Font size heuristic (larger text = heading)
                try:
                    for text_line in element:
                        for char in text_line:
                            if hasattr(char, 'size'):
                                run.font.size = Pt(max(8, min(char.size, 32)))
                                if char.size > 14:
                                    run.bold = True
                                break
                        break
                except Exception:
                    run.font.size = Pt(11)

                # Font family
                if is_arabic:
                    run.font.name = 'Arial'
                else:
                    run.font.name = 'Calibri'

    doc.save(docx_path)
    return {"success": True, "engine": "pdfminer-docx", "pageCount": page_count}


def fallback_libreoffice(pdf_path: str, out_dir: str) -> dict:
    result = subprocess.run(
        [
            "soffice", "--headless", "--norestore",
            "--infilter=writer_pdf_import",
            "--convert-to", "docx",
            "--outdir", out_dir,
            pdf_path,
        ],
        capture_output=True, text=True, timeout=120,
        env={**os.environ, "HOME": out_dir, "TMPDIR": out_dir, "SAL_USE_VCLPLUGIN": "svp"},
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice failed: {result.stderr[:300]}")
    return {"success": True, "engine": "libreoffice", "pageCount": 1}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"success": False, "error": "Usage: pdf_to_word.py input.pdf output.docx"}))
        sys.exit(1)

    pdf_path  = sys.argv[1]
    docx_path = sys.argv[2]
    out_dir   = str(Path(docx_path).parent)

    try:
        result = extract_with_pdfminer(pdf_path, docx_path)
        print(json.dumps(result))
    except Exception as e1:
        try:
            result = fallback_libreoffice(pdf_path, out_dir)
            # LibreOffice names output after input file
            lo_out = Path(out_dir) / (Path(pdf_path).stem + ".docx")
            if lo_out.exists() and str(lo_out) != docx_path:
                lo_out.rename(docx_path)
            print(json.dumps(result))
        except Exception as e2:
            print(json.dumps({
                "success": False,
                "error": f"pdfminer: {str(e1)[:200]} | libreoffice: {str(e2)[:200]}"
            }))
            sys.exit(1)

if __name__ == "__main__":
    main()
